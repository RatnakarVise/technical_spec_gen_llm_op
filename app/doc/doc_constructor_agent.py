import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    para = doc.add_heading(text, level=level)
    return para if para is not None else doc.add_paragraph(text, style="Heading1")
    # doc.add_heading(text, level=level)

def add_table(doc, colnames, rows):
    table = doc.add_table(rows=1, cols=len(colnames))
    table.style = "Light List"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(colnames):
        hdr_cells[i].text = str(col)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val) if val is not None else ""
    return table
def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar); 
    run._r.append(instrText); 
    run._r.append(fldChar2); 
    run._r.append(fldChar3)

def add_bookmark(paragraph, name, bid):
    """Create a collapsed bookmark on the given paragraph."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)


    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))


    # Insert start at beginning of paragraph, end at end
    paragraph._p.insert(0, start)
    paragraph._p.append(end)

def add_pageref_field(paragraph, bookmark_name):
    """Insert a PAGEREF field that will display the page of the bookmark."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); 
    fldBegin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText"); 
    instr.set(qn("xml:space"), "preserve"); 
    instr.text = f"PAGEREF {bookmark_name} \\h"

    fldSep = OxmlElement("w:fldChar"); 
    fldSep.set(qn("w:fldCharType"), "separate")

    result = OxmlElement("w:t")
    result.text = "1"   # dummy, will be replaced by Word after update

    fldEnd = OxmlElement("w:fldChar"); 
    fldEnd.set(qn("w:fldCharType"), "end")

    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(result)
    run._r.append(fldEnd)

def find_section_content(content_list, section_title):
    for sec in content_list:
        if sec.get('section_name', '').lower().strip() == section_title.lower().strip():
            return sec['content']
    return None

def find_all_table_like_chunks(text):
    """
    Splits text into possible tables (blocks with at least 2 lines containing pipes)
    and other text. Returns list of ('table', ...) and ('text', ...) chunks.
    """
    if not text or not text.strip():
        return []
    lines = text.splitlines()
    chunks = []
    buf = []
    in_table = False

    def flush(par, typ):
        s = "\n".join(par).strip()
        if s:
            chunks.append((typ, s))

    i = 0
    while i < len(lines):
        l = lines[i]
        if l.count('|') >= 1 and l.strip() and (i+1<len(lines) and lines[i+1].count('|')>=1):
            # Start of table block
            buf = [l]
            i += 1
            while i < len(lines) and lines[i].count('|') >= 1 and lines[i].strip():
                buf.append(lines[i])
                i += 1
            flush(buf, "table")
            buf = []
        else:
            if l.strip():
                flush([l], "text")
            i += 1
    return chunks

def parse_markdown_table(table_md):
    lines = [l.strip() for l in table_md.strip().splitlines() if l.strip()]
    if len(lines) < 2: return None, None
    # Classic markdown: at least header and |---|
    if not (lines[0].startswith('|') and lines[0].endswith('|')):
        return None, None
    rows = [[cell.strip() for cell in l.strip('|').split('|')] for l in lines]
    if len(rows) < 2: return None, None
    # Remove divider if present
    if re.match(r'^[-:\s|]+$', ''.join(lines[1])):
        del rows[1]
    colnames, data_rows = rows[0], rows[1:]
    if all(len(r) == len(colnames) for r in data_rows):
        return colnames, data_rows
    return None, None

def parse_github_style_table(table_md):
    lines = [l.strip() for l in table_md.strip().splitlines() if l.strip()]
    if len(lines) < 2: return None, None
    # 2nd line looks like github divider (---|---)
    if not re.match(r'^-+(\s*\|\s*-+)+$', lines[1]):
        return None, None
    colnames = [c.strip() for c in lines[0].split('|')]
    data_rows = [[c.strip() for c in l.split('|')] for l in lines[2:] if '|' in l]
    if all(len(r) == len(colnames) for r in data_rows):
        return colnames, data_rows
    return None, None

def parse_simple_pipe_table(table_md):
    """
    Fallback: Any rows with consistent number of pipes.
    """
    # Only rows that have | and not divider
    lines = [l.strip() for l in table_md.strip().splitlines() if '|' in l and not re.match(r'^-+(\s*\|\s*-+)+$', l)]
    if len(lines) < 2: return None, None
    rows = [[c.strip() for c in l.split('|')] for l in lines]
    ncols = len(rows[0])
    if all(len(row) == ncols for row in rows):
        colnames = rows[0]
        data_rows = rows[1:]
        return colnames, data_rows
    return None, None

def parse_any_delim_table(table_md):
    """
    Ultra-forgiving: Split on the *most common* delimiter if all rows same length.
    """
    lines = [l.strip() for l in table_md.strip().splitlines() if l.strip()]
    if len(lines) < 2: return None, None
    delimiters = ['|', '\t', '  +']  # pipe, tab, multi-space
    for delim in delimiters:
        try:
            if delim == '  +':
                rows = [re.split(r'  +', l) for l in lines]
            else:
                rows = [l.split(delim) for l in lines]
            ncols = len(rows[0])
            if all(len(r) == ncols for r in rows):
                return rows[0], rows[1:]
        except Exception:
            continue
    return None, None

def extract_arrow_flow(text):
    if not text:
        return ""
    for line in text.splitlines():
        line = line.strip("` ").strip()
        if "->" in line and not line.lower().startswith(('diagram', 'flow', 'legend', '#')):
            return line
    if "->" in text:
        return text.strip()
    return ""

def build_document(content, sections, flow_diagram_agent=None, diagram_dir="diagrams"):
    doc = Document()

    # Add main heading
    add_heading(doc, "Technical Specification Document", 0)
     # Precompute bookmark names (used by index + later applied to headings)
    
    add_heading(doc, "Index", 1)
    bookmark_names = [f"sec_{i+1}" for i in range(len(sections))]

    for i, section in enumerate(sections):
        title = section.get("title")
        bookmark_name = f"sec_{i+1}"

        # Insert paragraph for index entry
        p = doc.add_paragraph()
        run = p.add_run(f"{i+1}. {title}")
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run.add_tab()
        add_pageref_field(p, bookmark_name)

    # Set a right-aligned tab with dot leader for each index entry (looks like: Title ....... 3)
    right_pos = Inches(6.0)  # adjust if your page margins differ

    for i, section in enumerate(sections):
        title = section.get("title")
        bookmark_name = f"sec_{i+1}"
        header = f"{i+1}. {title}"
        # add_heading(doc, header, 1)
        p = add_heading(doc, header, 1)
        # Add bookmark for index
        # add_bookmark(p, f"sec_{i+1}", i+1)
        add_bookmark(p, bookmark_name, i+1)
        sec_content = find_section_content(content, title)

        # FLOW DIAGRAM SECTION HANDLING
        if title.strip().lower() == "flow diagram":
            diagram_img = None
            if flow_diagram_agent is not None and sec_content:
                try:
                    flow_line = extract_arrow_flow(sec_content)
                    if flow_line:
                        diagram_img = flow_diagram_agent.run(flow_line)  # <-- Returns BytesIO
                    else:
                        diagram_img = None
                except Exception as e:
                    print(f"Flow diagram agent error: {e}")
                    diagram_img = None
            if diagram_img:
                doc.add_picture(diagram_img, width=Inches(5.5))
            else:
                doc.add_paragraph("[Flow diagram not available]")
                continue  # Skip remaining processing for this section

        # Universal parsing for text+tables:
        chunks = find_all_table_like_chunks(sec_content or "")
        for typ, value in chunks:
            if typ == 'text':
                doc.add_paragraph(value)
            elif typ == 'table':
                colnames, rows = parse_markdown_table(value)
                if not (colnames and rows):
                    colnames, rows = parse_github_style_table(value)
                if not (colnames and rows):
                    colnames, rows = parse_simple_pipe_table(value)
                if not (colnames and rows):
                    colnames, rows = parse_any_delim_table(value)
                if colnames and rows:
                    add_table(doc, colnames, rows)
                else:
                    doc.add_paragraph(value)

    doc.add_paragraph("\nDocument generated by PWC AI-powered ABAP Tech Spec Assistant.")
    return doc